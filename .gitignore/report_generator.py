from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import uuid
from datetime import datetime
from templates.rtu_template import apply_rtu_margins


def add_cover_page(doc, data):

    title = doc.add_paragraph()

    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = title.add_run(
        data.project_title.upper()
    )

    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph()

    p = doc.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p.add_run(
        f"Submitted By\n{data.student_name}"
    )

    doc.add_paragraph()

    g = doc.add_paragraph()

    g.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    g.add_run(
        f"Guide: {data.guide_name}"
    )

    doc.add_page_break()

def add_certificate_page(doc, data):

    heading = doc.add_paragraph()

    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = heading.add_run("CERTIFICATE")

    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    certificate_text = f"""
    This is to certify that the project entitled
    "{data.project_title}"
    submitted by {data.student_name}
    has been carried out under my supervision.

    Guide: {data.guide_name}
    Department: {data.department}
    Session: {data.session}
    """

    doc.add_paragraph(certificate_text)

    doc.add_page_break()
def add_declaration_page(doc, data):

    heading = doc.add_paragraph()

    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = heading.add_run("DECLARATION")

    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    today = datetime.now().strftime("%d-%m-%Y")
    declaration_text = f"""
    I hereby declare that the project entitled
    "{data.project_title}"
    submitted in partial fulfillment of the requirements
    for the award of degree is my original work.

    The work presented in this report has not been
    submitted elsewhere for the award of any degree.

    Student Name: {data.student_name}
    Date: {today}
    """

    doc.add_paragraph(declaration_text)

    doc.add_paragraph()

    signature = doc.add_paragraph()

    signature.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    signature.add_run(data.student_name)

    doc.add_page_break()

def generate_report(data):

    doc = Document()

    apply_rtu_margins(doc)

    add_cover_page(doc, data)

    add_certificate_page(doc, data)

    add_declaration_page(doc, data)

    for section_data in data.sections:

        doc.add_heading(
            section_data.heading,
            level=1
        )

        for subsection in section_data.subsections:

            doc.add_heading(
                subsection.title,
                level=2
            )

            doc.add_paragraph(
                subsection.content
            )
    
    safe_title = re.sub(
    r'[^a-zA-Z0-9_ ]',
    '',
    data.project_title
    )
    safe_title = safe_title.replace(" ", "_")
    unique_id = uuid.uuid4().hex[:8]
    filename = f"{safe_title}_{unique_id}.docx"
    doc.save(filename)

    return filename