from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import uuid
from datetime import datetime
from templates.template_manager import apply_template
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

def add_page_numbers(doc):

    section = doc.sections[0]

    footer = section.footer

    paragraph = footer.paragraphs[0]

    paragraph.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def apply_text_formatting(paragraph, formatting):

    paragraph.paragraph_format.line_spacing = (
        formatting.line_spacing
    )

    for run in paragraph.runs:

        run.font.name = formatting.font_name

        run.font.size = Pt(
            formatting.font_size
        )

def add_toc_page(doc, data):

    heading = doc.add_paragraph()

    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = heading.add_run("TABLE OF CONTENTS")

    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    chapter_no = 1

    for section in data.sections:

        doc.add_paragraph(
            f"{chapter_no}. {section.heading}"
        )

        chapter_no += 1

    doc.add_page_break()

def add_cover_page(doc, data):

    # Logo
    try:
        doc.add_picture(
            "assets/college_logo.png",
            width=Inches(1.5)
        )

        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

    except:
        pass

    doc.add_paragraph()

    # Project Title
    title = doc.add_paragraph()

    title.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    run = title.add_run(
        data.project_title.upper()
    )

    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph()

    # Submitted By
    p = doc.add_paragraph()

    p.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    p.add_run(
        f"Submitted By\n{data.student_name}"
    )

    doc.add_paragraph()

    # Guide
    g = doc.add_paragraph()

    g.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    g.add_run(
        f"Guide: {data.guide_name}"
    )

    doc.add_page_break()

def add_certificate_page(doc, data):

    heading = doc.add_paragraph()

    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = heading.add_run("CERTIFICATE")

    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    certificate_text = f"""
    This is to certify that the project entitled
    "{data.project_title}"
    submitted by {data.student_name}
    has been carried out under my supervision.

    Guide: {data.guide_name}
    Department: {data.department}
    Session: {data.session}
    """

    doc.add_paragraph(certificate_text)

    doc.add_page_break()

def add_declaration_page(doc, data):

    heading = doc.add_paragraph()

    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = heading.add_run("DECLARATION")

    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    today = datetime.now().strftime("%d-%m-%Y")
    declaration_text = f"""
    I hereby declare that the project entitled
    "{data.project_title}"
    submitted in partial fulfillment of the requirements
    for the award of degree is my original work.
    The work presented in this report has not been
    submitted elsewhere for the award of any degree.
    Student Name: {data.student_name}
    Date: {today}
    """
    doc.add_paragraph(declaration_text)
    doc.add_paragraph()
    signature = doc.add_paragraph()
    signature.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    signature.add_run(data.student_name)

    doc.add_page_break()

def add_acknowledgement_page(doc, data):

    # Heading
    heading = doc.add_paragraph()

    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = heading.add_run("ACKNOWLEDGEMENT")

    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    # Acknowledgement Text
    ack_text = f"""
    I express my sincere gratitude to my project guide, {data.guide_name},
    for his valuable guidance, encouragement and continuous support throughout the development of this project.
    I would also like to thank the faculty members of the{data.department}
    for providing the necessary resources and knowledge that helped me complete this work successfully.
    I am grateful to my college for providing a conducive environment for learning and project development.
    Finally, I would like to thank my parents, family members and friends for their constant motivation, encouragement and support during the course of this project.

    Student Name:
    {data.student_name}
    """

    paragraph = doc.add_paragraph(ack_text)

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    doc.add_paragraph()

    # Signature
    sign = doc.add_paragraph()

    sign.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    sign.add_run(data.student_name)

    # Next Page
    doc.add_page_break()

def add_chapters(doc, data):

    chapter_no = 1

    for section_data in data.sections:

        chapter_heading = doc.add_paragraph()

        chapter_heading.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        run = chapter_heading.add_run(
            f"CHAPTER {chapter_no}"
        )

        run.bold = True
        run.font.size = Pt(16)

        title = doc.add_paragraph()

        title.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        title_run = title.add_run(
            section_data.heading
        )

        title_run.bold = True
        title_run.font.size = Pt(14)

        doc.add_paragraph()

        for subsection in section_data.subsections:

            heading = doc.add_heading(
            subsection.title,
             level=2
            )

            if data.template == "CUSTOM":

                apply_text_formatting(
                heading,
                data.formatting
            )

            paragraph = doc.add_paragraph(
            subsection.content
            )

            if data.template == "CUSTOM":

                apply_text_formatting(
                paragraph,
                data.formatting
            )

        doc.add_page_break()

        chapter_no += 1

def apply_custom_formatting(doc, formatting):

    section = doc.sections[0]

    section.left_margin = Cm(
        formatting.left_margin
    )

    section.right_margin = Cm(
        formatting.right_margin
    )

    section.top_margin = Cm(
        formatting.top_margin
    )

    section.bottom_margin = Cm(
        formatting.bottom_margin
    )

def generate_report(data):

    doc = Document()

    if data.template == "CUSTOM":

        apply_custom_formatting(
            doc,
            data.formatting
        )

    else:

        apply_template(
            doc,
            data.template
        )

    add_cover_page(doc, data)

    add_certificate_page(doc, data)

    add_declaration_page(doc, data)

    add_acknowledgement_page(doc, data)

    add_toc_page(doc, data)

    add_page_numbers(doc)

    add_chapters(doc, data)

    safe_title = re.sub(
        r'[^a-zA-Z0-9_ ]',
        '',
        data.project_title
    )

    safe_title = safe_title.replace(" ", "_")

    unique_id = uuid.uuid4().hex[:8]

    filename = f"{safe_title}_{unique_id}.docx"

    doc.save(filename)

    return filename